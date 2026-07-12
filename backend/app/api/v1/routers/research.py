"""
DSRA V2 — Research Sessions & Reports Router
=============================================
Endpoints for managing research session creation, listing, status checks,
running agent workflows as background tasks, and streaming SSE progress logs.
"""

import json
from typing import Any, Optional
import uuid
from uuid import UUID

from fastapi import APIRouter, BackgroundTasks, Depends, HTTPException, status, Response
from fastapi.responses import StreamingResponse
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.api.dependencies import get_current_user
from app.core.events import event_broker
from app.core.orchestrator import ResearchOrchestrator
from app.db.models.report import Report
from app.db.models.research_session import ResearchSession
from app.db.models.user import User
from app.db.repositories.report import ReportRepository
from app.db.repositories.research_session import ResearchSessionRepository
from app.db.repositories.user import UserRepository
from app.db.session import get_db_session
from app.core.security import decode_token
from app.schemas.api.reports import ReportListItemResponse, ReportResponse
from app.schemas.api.research import (
    ResearchSessionCreateRequest,
    ResearchSessionDetailResponse,
    ResearchSessionResponse,
    ResearchSessionStartResponse,
)
from app.schemas.common import SessionState, ReportStatus
from app.schemas.api.claims import ClaimResponse

router = APIRouter(tags=["Research Workspaces"])


@router.post("/sessions", response_model=ResearchSessionResponse, status_code=status.HTTP_201_CREATED)
async def create_session(
    request: ResearchSessionCreateRequest,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db_session),
) -> Any:
    """
    Create a new research workspace session owned by the authenticated user.
    """
    session_repo = ResearchSessionRepository(db)
    
    from datetime import datetime, timezone
    new_session = ResearchSession(
        id=uuid.uuid4(),
        user_id=current_user.id,
        topic=request.topic,
        state=SessionState.CREATED,
        depth=request.depth,
        max_iterations=request.max_iterations,
        iteration_count=0,
        created_at=datetime.now(timezone.utc),
        updated_at=datetime.now(timezone.utc)
    )
    
    await session_repo.create(new_session)
    await db.commit()
    await db.refresh(new_session)
    
    return ResearchSessionResponse(
        session_id=new_session.id,
        topic=new_session.topic,
        state=new_session.state,
        depth=new_session.depth,
        iteration_count=new_session.iteration_count,
        max_iterations=new_session.max_iterations,
        created_at=new_session.created_at,
        updated_at=new_session.updated_at,
        completed_at=new_session.completed_at
    )


@router.get("/sessions", response_model=list[ResearchSessionResponse])
async def list_sessions(
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db_session),
) -> Any:
    """
    List all research sessions belonging to the current authenticated user.
    """
    query = (
        select(ResearchSession)
        .where(ResearchSession.user_id == current_user.id)
        .order_by(ResearchSession.created_at.desc())
    )
    result = await db.execute(query)
    sessions = result.scalars().all()
    
    return [
        ResearchSessionResponse(
            session_id=s.id,
            topic=s.topic,
            state=s.state,
            depth=s.depth,
            iteration_count=s.iteration_count,
            max_iterations=s.max_iterations,
            created_at=s.created_at,
            updated_at=s.updated_at,
            completed_at=s.completed_at
        )
        for s in sessions
    ]


@router.get("/sessions/{session_id}", response_model=ResearchSessionDetailResponse)
async def get_session_details(
    session_id: UUID,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db_session),
) -> Any:
    """
    Fetch comprehensive details, source metrics, and timeline state of a specific session.
    """
    session_repo = ResearchSessionRepository(db)
    session = await session_repo.get_by_id_with_relationships(session_id)
    
    if not session:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Research session not found."
        )
        
    if session.user_id != current_user.id:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Access forbidden to this research workspace."
        )

    # Compile counts
    sources_count = len(session.sources)
    claims_count = len(session.claims)
    verified_claims_count = sum(1 for c in session.claims if str(c.status) in ("VERIFIED", "VerificationStatus.VERIFIED"))
    
    # Get active report if finalized
    report_id = None
    if session.reports:
        # Get the latest report
        latest_report = sorted(session.reports, key=lambda r: r.created_at, reverse=True)[0]
        report_id = latest_report.id

    # Format timeline execution log history
    agent_timeline = [
        {
            "agent": log.agent_name,
            "status": "FAILED" if log.error_message else "COMPLETED",
            "message": log.error_message or "",
            "timestamp": log.created_at.isoformat() if log.created_at else None
        }
        for log in sorted(session.execution_logs, key=lambda l: l.created_at)
    ]

    return ResearchSessionDetailResponse(
        session_id=session.id,
        topic=session.topic,
        state=session.state,
        depth=session.depth,
        iteration_count=session.iteration_count,
        max_iterations=session.max_iterations,
        created_at=session.created_at,
        updated_at=session.updated_at,
        completed_at=session.completed_at,
        sources_count=sources_count,
        claims_count=claims_count,
        verified_claims_count=verified_claims_count,
        report_id=report_id,
        agent_timeline=agent_timeline
    )


@router.get("/sessions/{session_id}/claims", response_model=list[ClaimResponse])
async def get_session_claims(
    session_id: UUID,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db_session),
) -> Any:
    """Fetch all claims for a specific session."""
    session_repo = ResearchSessionRepository(db)
    session = await session_repo.get_by_id_with_relationships(session_id)
    if not session or session.user_id != current_user.id:
        raise HTTPException(status_code=404, detail="Session not found")
    
    return [
        ClaimResponse(
            id=c.id,
            text=c.text,
            confidence=c.confidence,
            status=c.status,
            source_ids=c.source_ids
        ) for c in session.claims
    ]


@router.delete("/sessions/{session_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_session(
    session_id: UUID,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db_session),
) -> None:
    """
    Delete a research session and all associated database records (cascade).
    """
    session_repo = ResearchSessionRepository(db)
    session = await session_repo.get_by_id(session_id)
    
    if not session:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Research session not found."
        )
        
    if session.user_id != current_user.id:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Access forbidden."
        )

    await session_repo.delete(session_id)
    await db.commit()


async def _run_orchestrator_job(session_id: UUID) -> None:
    """Helper job to spin up orchestrator with a dedicated DB session connection."""
    from app.db.session import AsyncSessionLocal
    async with AsyncSessionLocal() as session:
        try:
            orchestrator = ResearchOrchestrator(session)
            await orchestrator.run_research_session(session_id)
        except Exception as e:
            # The orchestrator handles internal exceptions, but log anyway in case of engine errors
            import structlog
            structlog.get_logger(__name__).error("background_orchestrator_failed", error=str(e), session_id=str(session_id))


@router.post("/sessions/{session_id}/start", response_model=ResearchSessionStartResponse)
async def start_session_execution(
    session_id: UUID,
    background_tasks: BackgroundTasks,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db_session),
) -> Any:
    """
    Spawn the background research worker threads for a created session.
    """
    session_repo = ResearchSessionRepository(db)
    session = await session_repo.get_by_id(session_id)
    
    if not session:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Research session not found."
        )
        
    if session.user_id != current_user.id:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Access forbidden."
        )
        
    if session.state not in [SessionState.CREATED, SessionState.FAILED, SessionState.COMPLETED]:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Cannot start session in state {session.state}."
        )

    # Queue background task
    background_tasks.add_task(_run_orchestrator_job, session_id)
    
    # Return starting payload
    return ResearchSessionStartResponse(
        session_id=session_id,
        state=SessionState.PLANNING,
        stream_url=f"/api/v1/sessions/{session_id}/stream"
    )


@router.get("/sessions/{session_id}/stream")
async def stream_session_progress(
    session_id: UUID,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db_session),
) -> StreamingResponse:
    """
    Server-Sent Events (SSE) streaming channel returning real-time agent checkpoints.
    """
    session_repo = ResearchSessionRepository(db)
    session = await session_repo.get_by_id(session_id)
    
    if not session:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Research session not found."
        )
        
    if session.user_id != current_user.id:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Access forbidden."
        )

    async def event_generator():
        try:
            async for sse_event in event_broker.subscribe(session_id):
                # Standard SSE block formatting
                # Event lines separated by newlines
                yield f"event: {str(sse_event.event)}\n"
                yield f"data: {json.dumps(sse_event.data)}\n\n"
        except Exception as e:
            # Yield error event and abort
            yield f"event: error\n"
            yield f"data: {json.dumps({'message': str(e)})}\n\n"

    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no"
        }
    )


# ── Reports Endpoints ──────────────────────────────────────────────────

@router.get("/reports", response_model=list[ReportListItemResponse])
async def list_reports(
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db_session),
) -> Any:
    """
    List all compiled research report drafts for the current user.
    """
    query = (
        select(Report)
        .join(ResearchSession)
        .where(ResearchSession.user_id == current_user.id)
        .order_by(Report.created_at.desc())
    )
    result = await db.execute(query)
    reports = result.scalars().all()
    
    return [
        ReportListItemResponse(
            id=r.id,
            session_id=r.session_id,
            title=r.title,
            executive_summary_snippet=r.executive_summary[:200] + "...",
            status=ReportStatus(r.status),
            critique_score=r.critique_score,
            created_at=r.created_at
        )
        for r in reports
    ]


@router.get("/reports/{report_id}", response_model=ReportResponse)
async def get_report_details(
    report_id: UUID,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db_session),
) -> Any:
    """
    Fetch the complete structured report draft matching report_id.
    """
    report_repo = ReportRepository(db)
    report = await report_repo.get_by_id(report_id)
    
    if not report:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Report not found."
        )

    # Check ownership by loading the associated session
    query = select(ResearchSession).where(ResearchSession.id == report.session_id)
    res = await db.execute(query)
    session = res.scalar_one_or_none()
    
    if not session or session.user_id != current_user.id:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Access forbidden to this report."
        )

    # Convert database structured JSON fields to lists/schemas expected by response model
    # SQLAlchemy columns sections and references are JSON fields
    return ReportResponse(
        id=report.id,
        session_id=report.session_id,
        title=report.title,
        executive_summary=report.executive_summary,
        sections=report.sections,
        key_findings=report.key_findings,
        references=report.references,
        methodology_description=report.methodology,
        limitations=report.limitations,
        conclusion=report.conclusion,
        visualization=report.visualization,
        critique_score=report.critique_score,
        status=ReportStatus(report.status),
        export_paths=report.export_paths or {},
        created_at=report.created_at,
        finalized_at=report.finalized_at
    )

@router.get("/sessions/{session_id}/download/{format_ext}")
async def download_report_export(
    session_id: UUID,
    format_ext: str,
    token: str,
    db: AsyncSession = Depends(get_db_session)
):
    """
    Download the generated report in various formats (PDF, MD, HTML, JSON, DOCX, ZIP).
    Uses a token query parameter for authentication instead of Bearer header for easy browser downloading.
    """
    payload = decode_token(token)
    if not payload or payload.get("type") != "access":
        raise HTTPException(status_code=status.HTTP_401_UNAUTHORIZED, detail="Invalid token")
    
    user_repo = UserRepository(db)
    user_id_str = payload.get("sub")
    if not user_id_str:
        raise HTTPException(status_code=status.HTTP_401_UNAUTHORIZED, detail="Invalid token payload")
        
    try:
        user_id = uuid.UUID(user_id_str)
    except ValueError:
        raise HTTPException(status_code=status.HTTP_401_UNAUTHORIZED, detail="Invalid user ID format")
        
    user = await user_repo.get_by_id(user_id)
    if not user:
        raise HTTPException(status_code=status.HTTP_401_UNAUTHORIZED, detail="User not found")
        
    session_repo = ResearchSessionRepository(db)
    session = await session_repo.get_by_id(session_id)
    if not session or session.user_id != user.id:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Session not found")
        
    report_repo = ReportRepository(db)
    report = await report_repo.get_by_session_id(session_id)
    if not report:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Report not found")
        
    fmt = format_ext.lower()
    
    # Generate content based on format
    # In a real app, these would be cached in S3 or generated via dedicated libraries.
    # We generate a faithful mock representation here using the actual database report object.
    
    if fmt == "json":
        content = json.dumps({
            "title": report.title,
            "executive_summary": report.executive_summary,
            "sections": report.sections,
            "key_findings": report.key_findings,
            "methodology": report.methodology,
            "limitations": report.limitations,
            "conclusion": report.conclusion,
            "references": report.references,
        }, indent=2)
        return Response(content=content, media_type="application/json", headers={
            "Content-Disposition": f"attachment; filename=report_{session_id}.json"
        })
        
    # Generate Markdown base for MD, HTML, DOCX, PDF
    md_content = f"# {report.title}\n\n## Executive Summary\n{report.executive_summary}\n\n## Key Findings\n"
    for finding in report.key_findings:
        md_content += f"- {finding}\n"
    for sec in report.sections:
        md_content += f"\n## {sec.get('heading', 'Section')}\n{sec.get('content', '')}\n"
    md_content += f"\n## Methodology\n{report.methodology}\n"
    md_content += f"\n## Limitations\n{report.limitations}\n"
    md_content += f"\n## Conclusion\n{report.conclusion}\n"
    
    if fmt == "md":
        return Response(content=md_content, media_type="text/markdown", headers={
            "Content-Disposition": f"attachment; filename=report_{session_id}.md"
        })
        
    if fmt == "html":
        import markdown
        html_body = markdown.markdown(md_content)
        html_content = f"<html><head><title>{report.title}</title><style>body{{font-family:sans-serif;line-height:1.6;margin:40px auto;max-width:800px;color:#333;}}h1,h2{{color:#111;}}</style></head><body>{html_body}</body></html>"
        return Response(content=html_content, media_type="text/html", headers={
            "Content-Disposition": f"attachment; filename=report_{session_id}.html"
        })
        
    if fmt == "docx":
        import io
        import docx
        doc = docx.Document()
        doc.add_heading(report.title, 0)
        doc.add_heading('Executive Summary', level=1)
        doc.add_paragraph(report.executive_summary)
        doc.add_heading('Key Findings', level=1)
        for finding in report.key_findings:
            doc.add_paragraph(finding, style='List Bullet')
        for sec in report.sections:
            doc.add_heading(sec.get('heading', 'Section'), level=1)
            doc.add_paragraph(sec.get('content', ''))
            
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        return Response(content=doc_io.read(), media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={
            "Content-Disposition": f"attachment; filename=report_{session_id}.docx"
        })
        
    if fmt == "pdf":
        import io
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet
        
        pdf_io = io.BytesIO()
        doc = SimpleDocTemplate(pdf_io, pagesize=letter)
        styles = getSampleStyleSheet()
        Story = []
        
        Story.append(Paragraph(report.title, styles['Title']))
        Story.append(Spacer(1, 12))
        Story.append(Paragraph('Executive Summary', styles['Heading1']))
        Story.append(Paragraph(report.executive_summary, styles['Normal']))
        Story.append(Spacer(1, 12))
        
        for sec in report.sections:
            Story.append(Paragraph(sec.get('heading', 'Section'), styles['Heading1']))
            Story.append(Paragraph(sec.get('content', ''), styles['Normal']))
            Story.append(Spacer(1, 12))
            
        doc.build(Story)
        pdf_io.seek(0)
        return Response(content=pdf_io.read(), media_type="application/pdf", headers={
            "Content-Disposition": f"attachment; filename=report_{session_id}.pdf"
        })
        
    if fmt == "zip":
        import io
        import zipfile
        
        zip_io = io.BytesIO()
        with zipfile.ZipFile(zip_io, mode='w', compression=zipfile.ZIP_DEFLATED) as zf:
            zf.writestr(f"report_{session_id}.json", json.dumps(report.sections, indent=2))
            zf.writestr(f"report_{session_id}.md", md_content)
            
        zip_io.seek(0)
        return Response(content=zip_io.read(), media_type="application/zip", headers={
            "Content-Disposition": f"attachment; filename=report_package_{session_id}.zip"
        })
        
    raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Unsupported format")
